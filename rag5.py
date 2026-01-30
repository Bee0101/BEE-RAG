import streamlit as st
import google.generativeai as genai
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import OllamaEmbeddings
from langchain_core.documents import Document
import PyPDF2
import docx
import os
import pickle
import hashlib
from typing import List
from dotenv import load_dotenv

load_dotenv()

VECTOR_STORE_DIR = "vector_store"
VECTOR_STORE_PATH = os.path.join(VECTOR_STORE_DIR, "faiss_index")
METADATA_PATH = os.path.join(VECTOR_STORE_DIR, "metadata.pkl")

os.makedirs(VECTOR_STORE_DIR, exist_ok=True)

st.set_page_config(
    page_title="BeeSolver AI Assistant",
    layout="wide",
    initial_sidebar_state="expanded"
)

st.markdown("""
<style>
    /* Main chat container */
    .stChatMessage {
        padding: 0.75rem 1rem;
        margin-bottom: 0.75rem;
        border-radius: 15px;
        max-width: 75%;
        word-wrap: break-word;
    }
    
    /* User messages - right side with blue/purple gradient */
    div[data-testid="stChatMessage"]:has(div[data-testid="chatAvatarIcon-user"]) {
        flex-direction: row-reverse;
    }
    
    div[data-testid="stChatMessage"]:has(div[data-testid="chatAvatarIcon-user"]) .stChatMessage {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        margin-left: auto;
        margin-right: 0;
        border-radius: 20px 20px 5px 20px;
        box-shadow: 0 2px 5px rgba(102, 126, 234, 0.3);
    }
    
    /* Assistant messages - left side with light gray background */
    div[data-testid="stChatMessage"]:has(div[data-testid="chatAvatarIcon-assistant"]) .stChatMessage {
        background-color: #f1f3f4;
        color: #202124;
        margin-right: auto;
        margin-left: 0;
        border-radius: 20px 20px 20px 5px;
        box-shadow: 0 2px 5px rgba(0, 0, 0, 0.1);
    }
    
    /* Avatar styling */
    div[data-testid="chatAvatarIcon-user"] {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    }
    
    div[data-testid="chatAvatarIcon-assistant"] {
        background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
    }
    
    /* Chat input area */
    .stChatInputContainer {
        border-top: 2px solid #e8eaed;
        padding-top: 1rem;
        background-color: white;
    }
    
    /* Sidebar styling */
    section[data-testid="stSidebar"] {
        background-color: #f8f9fa;
        border-right: 1px solid #e8eaed;
    }
    
    section[data-testid="stSidebar"] .stButton button {
        width: 100%;
        border-radius: 8px;
        font-weight: 500;
    }
    
    /* File uploader styling */
    .stFileUploader {
        border-radius: 8px;
        border: 2px dashed #dadce0;
        padding: 1rem;
    }
    
    /* Success/Info/Warning messages */
    .stAlert {
        border-radius: 8px;
        border-left: 4px solid;
    }
    
    /* Title styling */
    h1 {
        color: #202124;
        font-weight: 600;
    }
    
    /* Stored document badges */
    .doc-badge {
        display: inline-block;
        padding: 0.25rem 0.75rem;
        background-color: #e8f0fe;
        color: #1967d2;
        border-radius: 12px;
        font-size: 0.85rem;
        margin: 0.25rem 0;
        font-weight: 500;
    }
</style>
""", unsafe_allow_html=True)
 
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []
if "vector_store" not in st.session_state:
    st.session_state.vector_store = None
if "documents_processed" not in st.session_state:
    st.session_state.documents_processed = False
if "document_metadata" not in st.session_state:
    st.session_state.document_metadata = []
if "gemini_api_key" not in st.session_state:
    st.session_state.gemini_api_key = os.getenv("GEMINI_API_KEY")

def get_file_hash(uploaded_file) -> str:
    """Generate SHA-256 hash of uploaded file for duplicate detection"""
    uploaded_file.seek(0)
    data = uploaded_file.read()
    uploaded_file.seek(0)
    return hashlib.sha256(data).hexdigest()

def load_vector_store():
    """Load existing vector store and metadata from disk"""
    try:
        if os.path.exists(VECTOR_STORE_PATH + ".faiss"):
            embeddings = OllamaEmbeddings(
                model="nomic-embed-text",
                base_url="http://localhost:11434"
            )
            store = FAISS.load_local(
                VECTOR_STORE_PATH,
                embeddings,
                allow_dangerous_deserialization=True
            )
            if os.path.exists(METADATA_PATH):
                with open(METADATA_PATH, "rb") as f:
                    st.session_state.document_metadata = pickle.load(f)
            return store
        return None
    except Exception as e:
        st.error(f" Error loading vector store: {str(e)}")
        return None

def save_vector_store(store):
    """Save vector store and metadata to disk"""
    store.save_local(VECTOR_STORE_PATH)
    with open(METADATA_PATH, "wb") as f:
        pickle.dump(st.session_state.document_metadata, f)

def delete_vector_store():
    """Delete all stored vectors and metadata"""
    for ext in [".faiss", ".pkl"]:
        path = VECTOR_STORE_PATH + ext
        if os.path.exists(path):
            os.remove(path)
    if os.path.exists(METADATA_PATH):
        os.remove(METADATA_PATH)

 
if st.session_state.vector_store is None:
    vs = load_vector_store()
    if vs:
        st.session_state.vector_store = vs
        st.session_state.documents_processed = True
 
def extract_text_from_pdf(f):
    """Extract text from PDF file"""
    reader = PyPDF2.PdfReader(f)
    return "\n".join(page.extract_text() or "" for page in reader.pages)

def extract_text_from_docx(f):
    """Extract text from DOCX file"""
    doc = docx.Document(f)
    return "\n".join(p.text for p in doc.paragraphs)

def extract_text_from_txt(f):
    """Extract text from TXT file"""
    return f.read().decode("utf-8")

def process_documents(uploaded_files) -> List[Document]:
    """
    Process uploaded documents with automatic duplicate detection.
    Documents already in vector store are automatically reused.
    """
    docs = []
    existing_hashes = {
        m["hash"] for m in st.session_state.document_metadata if "hash" in m
    }
    
    new_docs_count = 0
    reused_docs_count = 0

    for f in uploaded_files:
        file_hash = get_file_hash(f)
        
        # Check if document already exists in vector store
        if file_hash in existing_hashes:
            reused_docs_count += 1
            st.info(f" **'{f.name}'** already in database - automatically reusing stored data!")
            continue

        ext = f.name.split(".")[-1].lower()
        if ext == "pdf":
            text = extract_text_from_pdf(f)
        elif ext == "docx":
            text = extract_text_from_docx(f)
        elif ext == "txt":
            text = extract_text_from_txt(f)
        else:
            st.warning(f" Unsupported file type: {f.name}")
            continue

        if text.strip():
            docs.append(
                Document(
                    page_content=text,
                    metadata={"source": f.name, "hash": file_hash}
                )
            )
            st.session_state.document_metadata.append({
                "filename": f.name,
                "type": ext,
                "size": len(text),
                "hash": file_hash
            })
            new_docs_count += 1
    
     
    if reused_docs_count > 0:
        st.success(f"📚 {reused_docs_count} document(s) found in database - no re-upload needed!")
    if new_docs_count > 0:
        st.success(f"📄 {new_docs_count} new document(s) ready for processing")
    
    return docs

def create_vector_store(documents, existing_store=None):
    """Create or update vector store with new documents"""
    embeddings = OllamaEmbeddings(
        model="nomic-embed-text",
        base_url="http://localhost:11434"
    )

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )

    chunks = splitter.split_documents(documents)

    if not chunks:
        return existing_store, 0

    if existing_store:
        existing_store.add_documents(chunks)
        store = existing_store
    else:
        store = FAISS.from_documents(chunks, embeddings)

    save_vector_store(store)
    return store, len(chunks)
 
def get_relevant_context(query, store, k=3):
    """Retrieve relevant context from vector store"""
    docs = store.similarity_search(query, k=k)
    return "\n\n".join(d.page_content for d in docs)

def generate_response(query, context):
    """Generate AI response using Gemini with retrieved context"""
    genai.configure(api_key=st.session_state.gemini_api_key)
    model = genai.GenerativeModel("gemini-3-flash-preview")

    prompt = f"""
Context from documents:
{context}

User Question:
{query}

Please provide a clear, helpful, and accurate answer based on the context above. If the context doesn't contain relevant information, let the user know and provide a general response if possible.
"""
    return model.generate_content(prompt).text
 
# Sidebar for document management
with st.sidebar:
    st.title("📁 Document Management")
    st.markdown("---")

    uploaded_files = st.file_uploader(
        "Upload your documents",
        type=["pdf", "docx", "txt"],
        accept_multiple_files=True,
        help="Supported formats: PDF, DOCX, TXT. Duplicate documents are automatically detected!"
    )

    if uploaded_files and st.button(" Process & Save Documents", type="primary"):
        with st.spinner("Processing documents..."):
            docs = process_documents(uploaded_files)
            if docs:
                store, chunks = create_vector_store(
                    docs,
                    st.session_state.vector_store
                )
                st.session_state.vector_store = store
                st.session_state.documents_processed = True
                st.success(f"✨ Successfully processed {chunks} new chunks!")
            else:
                st.info(" All documents already in database - using existing data")

    st.markdown("---")
    
    if st.session_state.document_metadata:
        st.markdown("### 📚 Stored Documents")
        st.caption(f"Total: {len(st.session_state.document_metadata)} document(s)")
        
        for i, d in enumerate(st.session_state.document_metadata, 1):
            st.markdown(
                f'<div class="doc-badge">📄 {d["filename"]}</div>',
                unsafe_allow_html=True
            )
    else:
        st.info("No documents uploaded yet")

    st.markdown("---")
    
    if st.button("🗑️ Delete All Documents", type="secondary"):
        if st.session_state.document_metadata:
            delete_vector_store()
            st.session_state.vector_store = None
            st.session_state.document_metadata = []
            st.session_state.documents_processed = False
            st.success("All documents deleted!")
            st.rerun()


st.title("🐝 BeeSolver AI Assistant")
st.caption("Your intelligent document-based Q&A assistant")


for msg in st.session_state.chat_history:
    with st.chat_message(msg["role"]):
        st.write(msg["content"])


if user_input := st.chat_input("💬 How can i help you..."):
   
    st.session_state.chat_history.append(
        {"role": "user", "content": user_input}
    )
    
    with st.chat_message("user"):
        st.write(user_input)

     
    with st.chat_message("assistant"):
        with st.spinner("Thinking..."):
            context = ""
            if st.session_state.vector_store:
                context = get_relevant_context(
                    user_input, st.session_state.vector_store
                )

            response = generate_response(user_input, context)
            st.write(response)

     
    st.session_state.chat_history.append(
        {"role": "assistant", "content": response}
    )
    
    st.rerun()